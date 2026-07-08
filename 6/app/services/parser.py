"""Resume parsing service — PDF, DOCX, and OCR fallback."""
import re
import io
from pathlib import Path
from typing import Optional
from loguru import logger


def clean_text(text: str) -> str:
    """Normalize whitespace, remove control chars, keep readable content."""
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[ \t]{2,}', ' ', text)
    return text.strip()


def parse_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using pdfplumber with PyPDF2 fallback."""
    text = ""
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
            text = "\n".join(pages)
    except Exception as e:
        logger.warning(f"pdfplumber failed: {e}, trying PyPDF2")

    if not text.strip():
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            text = "\n".join(
                page.extract_text() or "" for page in reader.pages
            )
        except Exception as e:
            logger.warning(f"PyPDF2 failed: {e}")

    # OCR fallback for image-based PDFs
    if not text.strip():
        text = _ocr_pdf(file_bytes)

    return clean_text(text)


def _ocr_pdf(file_bytes: bytes) -> str:
    """Use pdf2image + pytesseract for scanned PDFs."""
    try:
        from pdf2image import convert_from_bytes
        import pytesseract
        images = convert_from_bytes(file_bytes, dpi=200)
        parts = [pytesseract.image_to_string(img, lang='eng') for img in images]
        return "\n".join(parts)
    except Exception as e:
        logger.error(f"OCR failed: {e}")
        return ""


def parse_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX including tables and headers."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        return clean_text("\n".join(parts))
    except Exception as e:
        logger.error(f"DOCX parsing failed: {e}")
        return ""


def parse_resume(file_bytes: bytes, filename: str) -> str:
    """Dispatch to the correct parser based on file extension."""
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return parse_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return parse_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
