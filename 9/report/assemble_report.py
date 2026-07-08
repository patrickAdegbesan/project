#!/usr/bin/env python3
"""
Assembles all chapter modules into a single .docx file.
Run: python3 assemble_report.py
Output: /home/path/code/project/9/report/SRMS_Project_Report.docx
"""

import sys
sys.path.insert(0, '/home/path/code/project/9/report')

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

from build_report_ch1 import CH1
from build_report_ch2 import CH2
from build_report_ch3 import CH3
from build_report_ch4 import CH4
from build_report_ch5 import CH5, REFERENCES

OUTPUT = '/home/path/code/project/9/report/SRMS_Project_Report.docx'

# ─── helpers ──────────────────────────────────────────────────────────────────

def set_para_format(para, double_space=True, first_indent=False):
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(6)
    if double_space:
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    if first_indent:
        pf.first_line_indent = Cm(1.27)

def apply_run_font(run, bold=False, italic=False, size=12):
    run.font.name  = 'Times New Roman'
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic

def set_cell_shading(cell, fill_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def add_page_number(doc):
    """Add page number to top-right header."""
    section = doc.sections[0]
    header  = section.header
    para    = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run()
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

# ─── document setup ───────────────────────────────────────────────────────────

def create_document():
    doc = Document()

    # Margins 2.54 cm all sides
    for section in doc.sections:
        section.top_margin    = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin   = Cm(2.54)
        section.right_margin  = Cm(2.54)

    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    add_page_number(doc)
    return doc

# ─── content renderers ────────────────────────────────────────────────────────

def render_title_page(doc):
    """Render the cover / title page."""
    for _ in range(4):
        p = doc.add_paragraph()
        set_para_format(p)

    def center_bold(text, size=12):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_format(p)
        r = p.add_run(text)
        apply_run_font(r, bold=True, size=size)
        return p

    center_bold("CALEB UNIVERSITY, LAGOS", 14)
    center_bold("FACULTY OF COMPUTING AND ENGINEERING SCIENCES", 12)
    center_bold("DEPARTMENT OF COMPUTER SCIENCE", 12)

    for _ in range(2):
        p = doc.add_paragraph()
        set_para_format(p)

    center_bold(
        "DESIGN AND IMPLEMENTATION OF A WEB-BASED STUDENT RECORD MANAGEMENT\n"
        "SYSTEM FOR NIGERIAN EDUCATIONAL INSTITUTIONS",
        14
    )

    for _ in range(2):
        p = doc.add_paragraph()
        set_para_format(p)

    center_bold("BY", 12)

    for _ in range(1):
        p = doc.add_paragraph()
        set_para_format(p)

    center_bold("[INSERT NAME]", 12)
    center_bold("[INSERT MATRIC NUMBER]", 12)

    for _ in range(2):
        p = doc.add_paragraph()
        set_para_format(p)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_format(p)
    r = p.add_run(
        "A Project Report Submitted to the Department of Computer Science,\n"
        "Faculty of Computing and Engineering Sciences, Caleb University, Lagos,\n"
        "in Partial Fulfilment of the Requirements for the Award of\n"
        "Bachelor of Science (B.Sc.) Degree in Computer Science"
    )
    apply_run_font(r, size=12)

    for _ in range(2):
        p = doc.add_paragraph()
        set_para_format(p)

    center_bold("Supervisor: [INSERT SUPERVISOR NAME]", 12)

    for _ in range(2):
        p = doc.add_paragraph()
        set_para_format(p)

    center_bold("JUNE 2026", 12)

    doc.add_page_break()


def render_block(doc, block):
    btype = block['type']

    if btype == 'chapter_title':
        doc.add_page_break()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_format(p, double_space=True)
        r = p.add_run(block['text'])
        apply_run_font(r, bold=True, size=14)

    elif btype == 'chapter_subtitle':
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_format(p)
        r = p.add_run(block['text'])
        apply_run_font(r, bold=True, size=14)
        doc.add_paragraph()  # spacing after

    elif btype == 'heading1':
        p = doc.add_paragraph()
        set_para_format(p)
        r = p.add_run(block['text'])
        apply_run_font(r, bold=True, size=12)

    elif btype == 'para':
        p = doc.add_paragraph()
        set_para_format(p, first_indent=True)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(block['text'])
        apply_run_font(r, size=12)

    elif btype == 'list':
        for i, item in enumerate(block['items'], 1):
            p = doc.add_paragraph()
            set_para_format(p)
            p.paragraph_format.left_indent  = Cm(1.27)
            p.paragraph_format.first_line_indent = Cm(-0.64)
            r = p.add_run(f"({chr(96+i)})\t{item}")
            apply_run_font(r, size=12)

    elif btype == 'definition':
        p = doc.add_paragraph()
        set_para_format(p, first_indent=True)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r1 = p.add_run(block['term'] + ': ')
        apply_run_font(r1, bold=True, italic=True, size=12)
        r2 = p.add_run(block['defn'])
        apply_run_font(r2, size=12)

    elif btype == 'table_caption':
        p = doc.add_paragraph()
        set_para_format(p, double_space=False)
        p.paragraph_format.space_before = Pt(12)
        r = p.add_run(block['text'])
        apply_run_font(r, bold=True, size=12)

    elif btype == 'table':
        headers = block['headers']
        rows    = block['rows']
        n_cols  = len(headers)
        tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
        tbl.style = 'Table Grid'

        # Header row
        hdr_cells = tbl.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            set_cell_shading(hdr_cells[i], '1a237e')
            for para in hdr_cells[i].paragraphs:
                for run in para.runs:
                    run.font.name  = 'Times New Roman'
                    run.font.size  = Pt(10)
                    run.font.bold  = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after  = Pt(2)

        # Data rows
        for r_idx, row_data in enumerate(rows):
            cells = tbl.rows[r_idx + 1].cells
            fill  = 'F5F5F5' if r_idx % 2 == 0 else 'FFFFFF'
            for c_idx, cell_text in enumerate(row_data):
                cells[c_idx].text = cell_text
                set_cell_shading(cells[c_idx], fill)
                for para in cells[c_idx].paragraphs:
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(9)
                    para.paragraph_format.space_before = Pt(2)
                    para.paragraph_format.space_after  = Pt(2)

        # Space after table
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)

    elif btype == 'mermaid_caption':
        p = doc.add_paragraph()
        set_para_format(p, double_space=False)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        r = p.add_run('[DIAGRAM PLACEHOLDER]')
        apply_run_font(r, italic=True, size=10)
        r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        # caption below placeholder
        p2 = doc.add_paragraph()
        set_para_format(p2, double_space=False)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(block['text'])
        apply_run_font(r2, bold=True, size=10)
        doc.add_paragraph()

    elif btype == 'mermaid':
        # The actual Mermaid code is stored in the .py files for reference;
        # in the docx we show the placeholder box. Mermaid code is appended as
        # a faint preformatted block so the author can render it separately.
        p = doc.add_paragraph()
        set_para_format(p, double_space=False)
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        r = p.add_run("(Mermaid diagram source available separately — render with Mermaid Live Editor and insert above)")
        apply_run_font(r, italic=True, size=9)
        r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    elif btype == 'screenshot':
        p_ph = doc.add_paragraph()
        p_ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_format(p_ph, double_space=False)
        p_ph.paragraph_format.space_before = Pt(12)
        r_ph = p_ph.add_run(f"[SCREENSHOT OF {block['label']} GOES HERE]")
        apply_run_font(r_ph, italic=True, size=11)
        r_ph.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

        if block.get('desc'):
            p_d = doc.add_paragraph()
            set_para_format(p_d, double_space=False)
            p_d.paragraph_format.left_indent = Cm(1.27)
            r_d = p_d.add_run(block['desc'])
            apply_run_font(r_d, italic=True, size=10)
            r_d.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_format(p_cap, double_space=False)
        r_cap = p_cap.add_run(block['caption'])
        apply_run_font(r_cap, bold=True, size=10)
        doc.add_paragraph()

    elif btype == 'code_caption':
        p = doc.add_paragraph()
        set_para_format(p, double_space=False)
        p.paragraph_format.space_before = Pt(10)
        r = p.add_run(block['text'])
        apply_run_font(r, bold=True, size=11)

    elif btype == 'code':
        p = doc.add_paragraph()
        set_para_format(p, double_space=False)
        p.paragraph_format.left_indent  = Cm(1.27)
        p.paragraph_format.right_indent = Cm(1.27)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        r = p.add_run(block['text'].strip())
        r.font.name = 'Courier New'
        r.font.size = Pt(8.5)
        # light grey box
        from docx.oxml import OxmlElement as OE
        pPr = p._p.get_or_add_pPr()
        shd = OE('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F0F0F0')
        pPr.append(shd)
        doc.add_paragraph()

    elif btype == 'reference':
        p = doc.add_paragraph()
        set_para_format(p, double_space=True)
        p.paragraph_format.left_indent        = Cm(1.27)
        p.paragraph_format.first_line_indent  = Cm(-1.27)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(block['text'])
        apply_run_font(r, size=12)


# ─── main assembly ────────────────────────────────────────────────────────────

def main():
    doc = create_document()
    render_title_page(doc)

    all_blocks = CH1 + CH2 + CH3 + CH4 + CH5 + REFERENCES

    for i, block in enumerate(all_blocks, 1):
        try:
            render_block(doc, block)
        except Exception as e:
            print(f"WARNING: block {i} ({block.get('type','?')}) failed: {e}")

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")

if __name__ == '__main__':
    main()
