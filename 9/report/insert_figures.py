#!/usr/bin/env python3
"""
Replaces [DIAGRAM PLACEHOLDER] paragraphs in the docx with actual PNG images.
"""
import os, re
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

DOCX_IN  = '/home/path/code/project/9/report/SRMS_Project_Report.docx'
DOCX_OUT = '/home/path/code/project/9/report/SRMS_Project_Report.docx'
FIG_DIR  = '/home/path/code/project/9/report/figures'

# Map caption text fragment → PNG filename
FIGURE_MAP = {
    'Figure 2.1': 'fig2_1_TAM.png',
    'Figure 2.2': 'fig2_2_evolution.png',
    'Figure 2.3': 'fig2_3_features.png',
    'Figure 3.1': 'fig3_1_usecase.png',
    'Figure 3.2': 'fig3_2_architecture.png',
    'Figure 3.3': 'fig3_3_erd.png',
    'Figure 3.4': 'fig3_4_dfd.png',
    'Figure 3.5': 'fig3_5_login.png',
    'Figure 3.6': 'fig3_6_registration.png',
    'Figure 3.7': 'fig3_7_grades.png',
    'Figure 3.8': 'fig3_8_fees.png',
    'Figure 3.9': 'fig3_9_reports.png',
    'Figure 3.10': 'fig3_10_schema.png',
    'Figure 3.11': 'fig3_11_deployment.png',
}

doc = Document(DOCX_IN)

# We need to find pairs: [DIAGRAM PLACEHOLDER] paragraph followed soon by
# a caption paragraph containing "Figure X.X: ..."
# Strategy: scan all paragraphs, when we find a placeholder, look ahead for
# the caption, determine which figure it is, replace the placeholder with image.

paras = doc.paragraphs

def find_figure_key(text):
    # Sort by length descending so "Figure 3.10" matches before "Figure 3.1"
    for key in sorted(FIGURE_MAP.keys(), key=len, reverse=True):
        # Use word-boundary style match: key must not be followed by a digit
        pattern = re.escape(key) + r'(?!\d)'
        if re.search(pattern, text):
            return key
    return None

# Build list of (index, paragraph) for placeholders
placeholder_indices = []
for i, p in enumerate(paras):
    if '[DIAGRAM PLACEHOLDER]' in p.text:
        placeholder_indices.append(i)

print(f'Found {len(placeholder_indices)} diagram placeholders')

replaced = 0
for ph_idx in placeholder_indices:
    # Look ahead up to 5 paragraphs for a caption
    fig_key = None
    caption_text = ''
    for j in range(ph_idx+1, min(ph_idx+6, len(paras))):
        txt = paras[j].text
        fk = find_figure_key(txt)
        if fk:
            fig_key = fk
            caption_text = txt
            break

    if not fig_key:
        print(f'  WARNING: no figure key found near placeholder at index {ph_idx}')
        continue

    img_path = os.path.join(FIG_DIR, FIGURE_MAP[fig_key])
    if not os.path.exists(img_path):
        print(f'  WARNING: image not found: {img_path}')
        continue

    # Replace the placeholder paragraph with the image
    p = paras[ph_idx]
    # Clear the paragraph
    for run in p.runs:
        run.text = ''
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add image run
    run = p.add_run()
    # Choose width based on figure type
    width = Inches(5.5) if '3.3' in fig_key or '3.1' in fig_key else Inches(5.8)
    run.add_picture(img_path, width=width)

    print(f'  Inserted {FIGURE_MAP[fig_key]} for {fig_key}')
    replaced += 1

doc.save(DOCX_OUT)
print(f'\nDone. {replaced} figures inserted. Saved to {DOCX_OUT}')
